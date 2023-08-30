
import tkinter as tk
from tkinter import ttk, filedialog, messagebox, Toplevel, Text, Scrollbar
from docx import Document
import os

class EnhancedUIApp:
    def __init__(self, root):
        self.template_paths = []
        self.entries = {}

        # Window title and size
        root.title("Генератор документов")
        root.geometry("650x700")

        # Button to select templates
        self.select_template_btn = ttk.Button(root, text="Загрузить шаблоны", command=self.select_templates)
        self.select_template_btn.grid(row=0, column=0, columnspan=2, pady=10)
        
        # Label to display the number of selected templates
        self.template_count_label = ttk.Label(root, text="Выбрано шаблонов: 0")
        self.template_count_label.grid(row=1, column=0, columnspan=2, pady=10)

        # Add new variable button and fields
        self.new_var_label = ttk.Label(root, text="Добавить переменную")
        self.new_var_label.grid(row=2, column=0, sticky=tk.W, padx=10, pady=5)
        
        self.new_var_name = ttk.Entry(root, width=20)
        self.new_var_name.grid(row=2, column=1, padx=10, pady=5)
        
        self.new_var_value = ttk.Entry(root, width=20)
        self.new_var_value.grid(row=2, column=2, padx=10, pady=5)

        self.add_var_btn = ttk.Button(root, text="Добавить", command=self.add_new_var)
        self.add_var_btn.grid(row=2, column=3, padx=10, pady=5)

        # Delete variable button
        self.del_var_btn = ttk.Button(root, text="Удалить", command=self.delete_var)
        self.del_var_btn.grid(row=2, column=4, padx=10, pady=5)

        # Frame for variables
        self.frame = ttk.Frame(root)
        self.frame.grid(row=3, column=0, columnspan=5, pady=10)

        # Scrollbar
        self.scrollbar = Scrollbar(self.frame)
        self.scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

        # Listbox to display variables
        self.listbox = tk.Listbox(self.frame, yscrollcommand=self.scrollbar.set, width=50, height=10)
        self.listbox.pack(side=tk.LEFT, fill=tk.BOTH)
        self.scrollbar.config(command=self.listbox.yview)

        # Dropdown to select file format
        self.file_format_label = ttk.Label(root, text="Выберите формат файла:")
        self.file_format_label.grid(row=4, column=0, padx=10, pady=5)
        self.file_format_var = tk.StringVar()
        self.file_format_dropdown = ttk.Combobox(root, textvariable=self.file_format_var, values=["docx", "txt"])
        self.file_format_dropdown.grid(row=4, column=1, padx=10, pady=5)
        self.file_format_dropdown.set("docx")

        # Create and place the generate and preview button
        self.generate_btn = ttk.Button(root, text="Сгенерировать и предпросмотреть документы", command=self.generate_and_preview_docs)
        self.generate_btn.grid(row=5, column=0, columnspan=5, pady=20)

        # Tooltips
        self._add_tooltip(self.select_template_btn, "Загрузите шаблоны документов в формате .docx")
        self._add_tooltip(self.new_var_name, "Введите имя новой переменной")
        self._add_tooltip(self.new_var_value, "Введите значение новой переменной")
        self._add_tooltip(self.add_var_btn, "Добавьте новую переменную в список")
        self._add_tooltip(self.del_var_btn, "Удалите выбранную переменную из списка")
        self._add_tooltip(self.generate_btn, "Сгенерируйте и предварительно просмотрите документы на основе шаблонов и переменных")
        self._add_tooltip(self.file_format_dropdown, "Выберите формат файла для сохранения: docx или txt")

    def _add_tooltip(self, widget, text):
        tool_tip = ToolTip(widget, text=text)

    def select_templates(self):
        filepaths = filedialog.askopenfilenames(title="Выберите шаблоны", filetypes=[("Word Documents", "*.docx")])
        if filepaths:
            self.template_paths.extend(filepaths)
            self.template_count_label.config(text=f"Выбрано шаблонов: {len(self.template_paths)}")

    def add_new_var(self):
        var_name = self.new_var_name.get()
        var_value = self.new_var_value.get()
        if var_name and var_value:
            self.entries[var_name] = var_value
            self.listbox.insert(tk.END, f"{var_name} = {var_value}")
            self.new_var_name.delete(0, tk.END)
            self.new_var_value.delete(0, tk.END)

    def delete_var(self):
        selected = self.listbox.curselection()
        if selected:
            var_name = self.listbox.get(selected).split(" = ")[0]
            del self.entries[var_name]
            self.listbox.delete(selected)

    def generate_and_preview_docs(self):
        generated_files = self._generate_documents(self.entries, self.template_paths)
        for file in generated_files:
            self.show_preview(file)

    def _read_document(self, filepath):
        doc = Document(filepath)
        return [paragraph.text for paragraph in doc.paragraphs]

    def _write_document(self, paragraphs, filepath, file_format="docx"):
        if file_format == "docx":
            doc = Document()
            for paragraph in paragraphs:
                doc.add_paragraph(paragraph)
            doc.save(filepath)
        elif file_format == "txt":
            with open(filepath, "w", encoding="utf-8") as file:
                file.write("\n".join(paragraphs))

    def _replace_variables_in_paragraph(self, paragraph, data):
        for key, value in data.items():
            paragraph = paragraph.replace(key, value)
        return paragraph

    def _generate_documents(self, data, template_paths):
        generated_files = []
        for template_path in template_paths:
            paragraphs = self._read_document(template_path)
            new_paragraphs = [self._replace_variables_in_paragraph(paragraph, data) for paragraph in paragraphs]
            output_path = f"generated_{os.path.basename(template_path)}"
            self._write_document(new_paragraphs, output_path)
            generated_files.append(output_path)
        return generated_files

    def show_preview(self, file_path):
        preview_window = Toplevel()
        preview_window.title(f"Предпросмотр {os.path.basename(file_path)}")
        text_widget = Text(preview_window, wrap=tk.WORD)
        text_widget.pack(fill=tk.BOTH, expand=True)
        with open(file_path, "r", encoding="utf-8") as file:
            text_widget.insert(tk.END, file.read())

def _add_tooltip(self, widget, text):
        tool_tip = ToolTip(widget, text=text)

class ToolTip:
    def __init__(self, widget, text):
        self.widget = widget
        self.text = text
        self.tooltip_window = None
        self.widget.bind("<Enter>", self.on_enter)
        self.widget.bind("<Leave>", self.on_leave)

    def on_enter(self, event=None):
        x, y, _, _ = self.widget.bbox("insert")
        x += self.widget.winfo_rootx() + 25
        y += self.widget.winfo_rooty() + 25
        self.tooltip_window = Toplevel(self.widget)
        self.tooltip_window.wm_overrideredirect(True)
        self.tooltip_window.wm_geometry(f"+{x}+{y}")
        label = tk.Label(self.tooltip_window, text=self.text, background="yellow", relief="solid", borderwidth=1)
        label.pack()

    def on_leave(self, event=None):
        if self.tooltip_window:
            self.tooltip_window.destroy()
            self.tooltip_window = None

    def show_tooltip(self):
        self.on_enter()

    def hide_tooltip(self):
        self.on_leave()

if __name__ == "__main__":
    root = tk.Tk()
    app = EnhancedUIApp(root)
    root.mainloop()
